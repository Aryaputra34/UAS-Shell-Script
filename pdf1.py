from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import win32com.client as win32
import os

# Tentukan nama file yang ingin digunakan
nama_file = 'INGFO SERVER.docx'
nomor = 1
file_path = nama_file

# Periksa apakah nama file sudah ada, jika ya, tambahkan nomor unik
while os.path.exists(file_path):
    nomor += 1
    nama_file_tanpa_ekstensi, ekstensi = os.path.splitext(nama_file)
    file_path = f'{nama_file_tanpa_ekstensi}_{nomor}{ekstensi}'

# Membuat objek Document
doc = Document()
gambar_path = 'download.jpg'
paragraf_gambar = doc.add_paragraph()
paragraf_gambar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_gambar = paragraf_gambar.add_run()
gambar = run_gambar.add_picture(gambar_path, width=Inches(1.5), height=Inches(1.5))

# Tambahkan paragraf
paragraf = doc.add_paragraph()
run = paragraf.add_run('Pemberitahuan Kondisi Server')
font = run.font
font.size = Pt(14)  # Mengatur ukuran font menjadi 14 Pt

# Mengatur paragraf menjadi rata tengah
paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

paragraf = doc.add_paragraph()
run = paragraf.add_run('Politeknik Negeri Jakarta')
font = run.font
font.size = Pt(14)  # Mengatur ukuran font menjadi 14 Pt

# Mengatur paragraf menjadi rata tengah
paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

sukses="berhasil ping dan nginx"
nginx="berhasil nginx gagal ping"
ping="berhasil ping gagal nginx"
down="server mati, tidak bisa ping ke website dan web server down"
kondisi1 = "berhasil ping dan nginx"

if kondisi1 == sukses:
    doc.add_paragraph('Server (PING) : [✓] OK', style='List Bullet')
    doc.add_paragraph('Web Server (NGINX) : [✓] OK', style='List Bullet')

elif kondisi1 == nginx:
    doc.add_paragraph('Server (PING): [✗] ERROR/OFF', style='List Bullet')
    doc.add_paragraph('Web Server (NGINX) : [✓] OK', style='List Bullet')
elif kondisi1 == ping:
    doc.add_paragraph('Server (PING) : [✓] OK', style='List Bullet')
    doc.add_paragraph('Web Server (NGINX) : [✗] ERROR/OFF', style='List Bullet')
elif kondisi1 == down:
    doc.add_paragraph('Server (PING): [✗] ERROR/OFF', style='List Bullet')
    doc.add_paragraph('Web Server (NGINX) : [✗] ERROR/OFF', style='List Bullet')
else :
    doc.add_paragraph('ERROR CODE!!!')

doc.save(file_path)
print(f'Dokumen disimpan sebagai: {file_path}')

nama_file_word = 'INGFO SERVER.docx'

# Tentukan nama file output PDF
nama_file_pdf = 'INGFO SERVER.pdf'

# Periksa apakah nama file PDF sudah ada, jika ya, tambahkan nomor unik
nomor = 1
nama_file_pdf_tanpa_ekstensi, ekstensi = os.path.splitext(nama_file_pdf)
file_path_pdf = nama_file_pdf
while os.path.exists(file_path_pdf):
    nomor += 1
    file_path_pdf = f'{nama_file_pdf_tanpa_ekstensi}_{nomor}{ekstensi}'

# Konversi dokumen Word ke PDF
word = win32.gencache.EnsureDispatch('Word.Application')
doc = word.Documents.Open(os.path.abspath(nama_file_word))
doc.SaveAs(os.path.abspath(file_path_pdf), FileFormat=17)  # FileFormat 17 untuk PDF
doc.Close()
word.Quit()

print(f'Dokumen PDF berhasil dibuat: {file_path_pdf}')
os.remove(file_path)
